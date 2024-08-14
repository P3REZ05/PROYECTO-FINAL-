from flask import Flask, render_template, request, redirect, url_for, flash, jsonify, session, send_file
import os
from datetime import date
import mysql.connector
from database import get_db_cursor, db
import database as db
from database import database
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import io
import pandas as pd
from docx import Document



# Configuración
template_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), 'templates'))
app = Flask(__name__, template_folder=template_dir)
app.secret_key = 'your_secret_key'

# Conexión a la base de datos



# Ruta para el formulario de inicio de sesión
@app.route('/')
def index():
    return render_template('index.html')

@app.route('/login', methods=['POST'])
def login():
    if request.method == 'POST':

        cursor = db.db.cursor(dictionary=True)
        username = request.form['username']
        password = request.form['password']
        rol = request.form['rol']

        # Utilizamos db.db para obtener la conexión

    if rol == 'empleado':
        cursor.execute('SELECT * FROM usuarios WHERE Nombre = %s AND cedula = %s', (username, password))
    else:
        cursor.execute('SELECT * FROM admin WHERE nombre = %s AND cc = %s', (username, password))

    user = cursor.fetchone()

    if user:
        session['id_usuario'] = user['ID_usuario'] if rol == 'empleado' else user['cc']
        if rol == 'empleado':
            return redirect(url_for('ventas'))
        else:
            return redirect(url_for('dashboard'))

    flash('Invalid username or password')
    return redirect(url_for('index'))




# Ruta para la gestión de usuarios
@app.route('/usuarios')
def usuarios():


    cursor = db.db.cursor(dictionary=True)


    cursor = database.cursor()
    cursor.execute("SELECT * FROM usuarios")
    myresult = cursor.fetchall()
    
    # Convertir datos a diccionario
    insertObject = []
    columnNames = [column[0] for column in cursor.description]
    
    for record in myresult:
        insertObject.append(dict(zip(columnNames, record)))
    
    cursor.close()
    return render_template('usuarios.html', data=insertObject, title='usuarios')

# Ruta para guardar usuarios en la base de datos
@app.route('/usuarios', methods=['POST'])
def guardar_usuario():
    ID_usuario = request.form['ID_usuario']
    Nombre = request.form['Nombre']
    cedula = request.form['Cedula']
    fecha_registro = request.form['fecha_registro']

    if ID_usuario and Nombre and cedula and fecha_registro:
        try:
            cursor = database.cursor()
            sql = "INSERT INTO usuarios (ID_usuario, Nombre, cedula, fecha_registro) VALUES (%s, %s, %s, %s)"
            data = (ID_usuario, Nombre, cedula, fecha_registro)
            cursor.execute(sql, data)
            database.commit()
        except mysql.connector.errors.IntegrityError:
            flash('Error: El ID de usuario ya existe.', 'error')
        except mysql.connector.errors.DataError:
            flash('Error: El valor de la cédula es demasiado largo.', 'error')

    return redirect(url_for('usuarios'))

# Ruta para eliminar usuarios
@app.route('/delete/<string:id>')
def delete(id):
    cursor = database.cursor()
    sql = "DELETE FROM usuarios WHERE ID_usuario = %s"
    data = (id,)
    cursor.execute(sql, data)
    database.commit()
    return redirect(url_for('usuarios'))

# Ruta para editar usuarios 
@app.route('/edit/<string:id>', methods=['POST'])
def edit(id):
    ID_usuario = request.form['ID_usuario']
    Nombre = request.form['Nombre']
    cedula = request.form['Cedula']
    fecha_registro = request.form['fecha_registro']

    if ID_usuario and Nombre and cedula and fecha_registro:
        try:
            cursor = database.cursor()
            sql = "UPDATE usuarios SET ID_usuario = %s, Nombre = %s, cedula = %s, fecha_registro = %s WHERE ID_usuario = %s"
            data = (ID_usuario, Nombre, cedula, fecha_registro, id)
            cursor.execute(sql, data)
            database.commit()
        except mysql.connector.errors.IntegrityError:
            flash('Error: El ID de usuario ya existe.', 'error')
        except mysql.connector.errors.DataError:
            flash('Error: El valor de la cédula es demasiado largo.', 'error')

    return redirect(url_for('usuarios'))

# Ruta para mostrar la página de proveedores
# Ruta para mostrar la página de proveedores
@app.route('/proveedores')
def mostrar_proveedores():
    try:
        cursor = database.cursor(dictionary=True)
        cursor.execute("SELECT * FROM proveedores")
        myresult = cursor.fetchall()
        cursor.close()
        
        return render_template('proveedores.html', data=myresult)
    
    except mysql.connector.Error as err:
        print(f"Error al obtener proveedores de la base de datos: {err}")
        return render_template('proveedores.html', data=[], title='proveedores')

# Ruta para guardar proveedores en la base de datos
@app.route('/proveedores/guardar', methods=['POST'])
def guardar_proveedor():
    ID_proveedor = request.form['ID_proveedor']
    Nombre_empresa = request.form['Nombre_empresa']
    Nombre_representante_legal = request.form['Nombre_representante_legal']
    tipo_proveedor = request.form['tipo_proveedor']
    direccion = request.form['direccion']
    telefono = request.form['telefono']
    email = request.form['email']
    URL = request.form['URL']

    if ID_proveedor and Nombre_empresa and Nombre_representante_legal and tipo_proveedor and direccion and telefono and email and URL:
        try:
            cursor = database.cursor()
            sql = "INSERT INTO proveedores (ID_proveedor, Nombre_empresa, Nombre_representante_legal, tipo_proveedor, direccion, telefono, email, URL) VALUES (%s, %s, %s, %s, %s, %s, %s, %s)"
            data = (ID_proveedor, Nombre_empresa, Nombre_representante_legal, tipo_proveedor, direccion, telefono, email, URL)
            cursor.execute(sql, data)
            database.commit()
        except mysql.connector.errors.IntegrityError:
            flash('Error: El ID de proveedor ya existe.', 'error')
        except mysql.connector.errors.DataError:
            flash('Error: Uno de los valores es demasiado largo.', 'error')

    return redirect(url_for('mostrar_proveedores'))

# Ruta para eliminar proveedores
@app.route('/proveedores/delete/<string:id>')
def eliminar_proveedor(id):
    try:
        cursor = database.cursor()
        sql = "DELETE FROM proveedores WHERE ID_proveedor = %s"
        data = (id,)
        cursor.execute(sql, data)
        database.commit()
    except mysql.connector.Error as err:
        print(f"Error al eliminar proveedor: {err}")
    finally:
        cursor.close()
    
    return redirect(url_for('mostrar_proveedores'))

# Ruta para editar proveedores 
@app.route('/proveedores/editar/<string:id>', methods=['POST'])
def editar_proveedor(id):
    ID_proveedor = request.form['ID_proveedor']
    Nombre_empresa = request.form['Nombre_empresa']
    Nombre_representante_legal = request.form['Nombre_representante_legal']
    tipo_proveedor = request.form['tipo_proveedor']
    direccion = request.form['direccion']
    telefono = request.form['telefono']
    email = request.form['email']
    URL = request.form['URL']

    if ID_proveedor and Nombre_empresa and Nombre_representante_legal and tipo_proveedor and direccion and telefono and email and URL:
        try:
            cursor = database.cursor()
            sql = "UPDATE proveedores SET ID_proveedor = %s, Nombre_empresa = %s, Nombre_representante_legal = %s, tipo_proveedor = %s, direccion = %s, telefono = %s, email = %s, URL = %s WHERE ID_proveedor = %s"
            data = (ID_proveedor, Nombre_empresa, Nombre_representante_legal, tipo_proveedor, direccion, telefono, email, URL, id)
            cursor.execute(sql, data)
            database.commit()
        except mysql.connector.errors.IntegrityError:
            flash('Error: El ID de proveedor ya existe.', 'error')
        except mysql.connector.errors.DataError:
            flash('Error: Uno de los valores es demasiado largo.', 'error')

    return redirect(url_for('mostrar_proveedores'))

# Ruta para mostrar la página de inventario
# Ruta principal del inventario
@app.route('/inventario')
def inventario():
    # if 'id_usuario' not in session:
    #     return redirect(url_for('index'))

    cursor = get_db_cursor()

    cursor.execute("SELECT * FROM producto")
    productos = cursor.fetchall()

    cursor.execute("SELECT * FROM proveedores")
    proveedores = cursor.fetchall()

    cursor.execute("SELECT * FROM categoria_producto")
    categorias = cursor.fetchall()

    cursor.close()

    return render_template('inventario.html', data=productos, proveedores=proveedores, categorias=categorias, title='inventario')

# Ruta para filtrar productos
@app.route('/filtrar_productos', methods=['POST'])
def filtrar_productos():
    proveedor = request.form.get('proveedor')
    precio = request.form.get('precio')
    cantidad = request.form.get('cantidad')

    cursor = get_db_cursor()

    # Construir la consulta SQL base
    sql = "SELECT * FROM producto WHERE 1=1"

    # Aplicar filtros según los parámetros recibidos
    if proveedor:
        sql += f" AND ID_proveedor = {proveedor}"
    if precio:
        sql += f" AND valor_producto < {precio}"
    if cantidad:
        sql += f" AND cantidad_inicial > {cantidad}"

    cursor.execute(sql)
    productos_filtrados = cursor.fetchall()
    cursor.close()

    # Obtener lista de proveedores para mostrar en la página
    cursor = get_db_cursor()
    cursor.execute("SELECT * FROM proveedores")
    proveedores = cursor.fetchall()
    cursor.close()
    
    return render_template('inventario.html', data=productos_filtrados, proveedores=proveedores)

# Función para agregar productos
@app.route('/agregar_producto', methods=['POST'])
def agregar_producto():
    ID_producto = request.form['ID_producto']
    nombre_producto = request.form['nombre_producto']
    descripcion_producto = request.form['descripcion_producto']
    valor_producto = request.form['valor_producto']
    stock = request.form['stock']
    ID_proveedor = request.form['ID_proveedor']
    ID_categoria_producto = request.form['ID_categoria_producto']

    if ID_producto and nombre_producto and descripcion_producto and valor_producto and stock and ID_proveedor and ID_categoria_producto:
        try:
            cursor = get_db_cursor()
            sql = "INSERT INTO producto (ID_producto, nombre_producto, descripcion_producto, valor_producto, stock, ID_proveedor, ID_categoria_producto) VALUES (%s, %s, %s, %s, %s, %s, %s)"
            data = (ID_producto, nombre_producto, descripcion_producto, valor_producto, stock, ID_proveedor, ID_categoria_producto)
            cursor.execute(sql, data)
            db.db.commit()
            cursor.close()
            flash('Producto agregado correctamente', 'success')
        except mysql.connector.Error as err:
            flash(f'Error al agregar producto: {err}', 'error')

    return redirect(url_for('inventario'))


#EJEMPLO CATEGORIA 

@app.route('/agregar_categoria', methods=['POST'])
def agregar_categoria():
    nombre_categoria = request.form['nombre_categoria']

    if nombre_categoria:
        try:
            cursor = database.cursor()
            sql = "INSERT INTO categoria_producto (nombre_categoria) VALUES (%s)"
            data = (nombre_categoria,)
            cursor.execute(sql, data)
            database.commit()
            cursor.close()
            flash('Categoría agregada correctamente', 'success')
            return redirect(url_for('inventario'))  # Return redirect here
        except mysql.connector.Error as err:
            flash(f'Error al agregar categoría: {err}', 'error')
            return redirect(url_for('inventario'))  # Return redirect here
    return redirect(url_for('inventario'))  # Return redirect here


# Ruta para editar categorías
@app.route('/editar_categoria', methods=['POST'])
def editar_categoria():
    ID_categoria_producto = request.form['ID_categoria_producto']
    nuevo_nombre_categoria = request.form['nuevo_nombre_categoria']

    if ID_categoria_producto and nuevo_nombre_categoria:
        try:
            cursor = database.cursor()
            sql = "UPDATE categoria_producto SET nombre_categoria = %s WHERE ID_categoria_producto = %s"
            data = (nuevo_nombre_categoria, ID_categoria_producto)
            cursor.execute(sql, data)
            database.commit()
            cursor.close()
            flash('Categoría actualizada correctamente', 'success')
        except mysql.connector.Error as err:
            flash(f'Error al actualizar categoría: {err}', 'error')

    return redirect(url_for('inventario'))





#FINAL EJEMPLO CATEGORIA



# Ruta para eliminar productos
@app.route('/delete_producto/<string:id>')
def delete_producto(id):
    cursor = get_db_cursor()
    sql = "DELETE FROM producto WHERE ID_producto = %s"
    data = (id,)
    cursor.execute(sql, data)
    database.commit()  # Cambia aquí de `db` a `database`
    cursor.close()
    flash('Producto eliminado correctamente', 'success')
    return redirect(url_for('inventario'))


# Ruta para editar productos 
@app.route('/editar_producto/<string:id>', methods=['POST'])
def editar_producto_id(id):
    ID_producto = request.form['ID_producto']
    nombre_producto = request.form['nombre_producto']
    descripcion_producto = request.form['descripcion_producto']
    valor_producto = request.form['valor_producto']
    cantidad_inicial = request.form['cantidad_inicial']
    ID_proveedor = request.form['ID_proveedor']

    if ID_producto and nombre_producto and descripcion_producto and valor_producto and cantidad_inicial and ID_proveedor:
        try:
            cursor = get_db_cursor()  # Obtiene el cursor desde la conexión global
            sql = """
            UPDATE producto 
            SET nombre_producto=%s, descripcion_producto=%s, valor_producto=%s, cantidad_inicial=%s, ID_proveedor=%s
            WHERE ID_producto=%s
            """
            cursor.execute(sql, (nombre_producto, descripcion_producto, valor_producto, cantidad_inicial, ID_proveedor, ID_producto))
            db.commit()  # Usa la conexión global para hacer commit
            return redirect(url_for('inventario'))
        except Exception as e:
            print(f"Error: {e}")
            db.rollback()  # Usa la conexión global para hacer rollback
        finally:
            cursor.close()  # Cierra el cursor, no la conexión
    return redirect(url_for('inventario'))




# Ruta para eliminar una categoría
@app.route('/delete_category', methods=['POST'])
def delete_categoria():
    ID_categoria_producto = request.form['ID_categoria_producto']

    if ID_categoria_producto:
        try:
            cursor = get_db_cursor()
            sql = "DELETE FROM categoria_producto WHERE ID_categoria_producto = %s"
            cursor.execute(sql, (ID_categoria_producto,))
            cursor._connection.commit()
            cursor.close()
            flash('Categoría eliminada correctamente', 'success')
        except mysql.connector.Error as err:
            flash(f'Error al eliminar categoría: {err}', 'error')

    return redirect(url_for('inventario'))




# Ruta para el dashboard
@app.route('/dashboard')
def dashboard():
    # Consulta para obtener la cantidad de productos
    cursor = database.cursor()
    cursor.execute('SELECT Nombre_producto, stock FROM producto ORDER BY stock DESC LIMIT 7')
    data_productos = cursor.fetchall()
    nombres_productos = [producto[0] for producto in data_productos]
    stock = [producto[1] for producto in data_productos]


    cursor.execute('''
    SELECT p.Nombre_producto, dd.descripcion as Motivo_devolucion
    FROM detalle_devolucion dd
    JOIN producto p ON dd.ID_producto = p.ID_producto 
    ORDER BY RAND()
    LIMIT 4
''')
    productos_devueltos = cursor.fetchall()


    # Consulta para obtener las ventas en una fecha específica
    fecha_especifica = '2024-07-15'  # Ajusta a la fecha deseada
    cursor.execute('''
        SELECT SUM(dv.cantidad) as total_ventas
        FROM detalle_ventas dv
        JOIN ventas v ON dv.ID_venta = v.ID_venta
        WHERE v.fecha_venta = %s
    ''', (fecha_especifica,))
    total_ventas = cursor.fetchone()[0] or 0

    # Consulta para obtener los productos más vendidos
    cursor.execute('''
        SELECT p.Nombre_producto, SUM(dv.cantidad) as Cantidad_vendida
        FROM detalle_ventas dv
        JOIN producto p ON dv.ID_producto = p.ID_producto
        GROUP BY p.Nombre_producto
        ORDER BY Cantidad_vendida DESC
        LIMIT 7
    ''')
    productos_mas_vendidos = cursor.fetchall()

    cursor.close()

    return render_template('dashboard.html', 
                        nombres_productos=nombres_productos, 
                        stock=stock,
                        productos_devueltos=productos_devueltos,
                        fecha_especifica=fecha_especifica,
                        total_ventas=total_ventas,
                        productos_mas_vendidos=productos_mas_vendidos, title='Dashboard')



# Ruta para el dashboard de ventas
@app.route('/dashboard_ventas')
def dashboard_ventas():

    # Establecer conexión con la base de datos
    cursor = database.cursor(dictionary=True)

    # Consulta para obtener los trabajadores con más ventas
    cursor.execute('''
        SELECT u.Nombre AS nombre, COUNT(v.ID_venta) AS ventas
        FROM ventas v
        JOIN usuarios u ON v.ID_usuario = u.ID_usuario
        GROUP BY u.Nombre
        ORDER BY ventas DESC
        LIMIT 7
    ''')
    trabajadores = cursor.fetchall()

    # Consulta para obtener los productos con bajo stock
    cursor.execute('''
        SELECT Nombre_producto AS nombre, stock AS stock
        FROM producto
        ORDER BY stock ASC
        LIMIT 7
    ''')
    productos = cursor.fetchall()

    # Consulta para obtener los productos con precios más bajos
    cursor.execute('''
        SELECT Nombre_producto AS nombre, valor_producto AS precio
        FROM producto
        ORDER BY valor_producto ASC
        LIMIT 10
    ''')
    productos_precios = cursor.fetchall()

    # Cerrar cursor después de ejecutar las consultas
    cursor.close()

    # Renderizar la plantilla HTML con los datos obtenidos
    return render_template('dashboard_ventas.html', 
                        trabajadores=trabajadores, 
                        productos=productos, 
                        productos_precios=productos_precios, title='Dashboard Ventas')



#DANILOOOOO COMIENZO


@app.route('/ventas', methods=['GET', 'POST'])
def ventas():

    
    cursor = db.db.cursor(dictionary=True)
    

    categorias = []
    productos = []
    selected_categoria_id = request.form.get('categoria_id', None)
    search_query = request.form.get('search_query', '')

    venta_id = session.get('venta_id')

    if request.method == 'POST':
        if 'id_producto' in request.form and 'cantidad' in request.form:
            id_producto = request.form['id_producto']
            cantidad = int(request.form['cantidad'])

            cursor.execute('SELECT * FROM producto WHERE ID_producto = %s', (id_producto,))
            producto = cursor.fetchone()

            if producto and producto['stock'] >= cantidad:
                try:
                    if not venta_id:
                        # Crear una nueva venta
                        cursor.execute('INSERT INTO ventas (fecha_venta, ID_usuario) VALUES (NOW(), %s)', (session['id_usuario'],))
                        db.db.commit()
                        venta_id = cursor.lastrowid
                        session['venta_id'] = venta_id

                    # Insertar detalle de venta
                    cursor.execute('INSERT INTO detalle_ventas (ID_venta, ID_producto, cantidad, valor_venta_producto) VALUES (%s, %s, %s, %s)', 
                                (venta_id, id_producto, cantidad, producto['valor_producto']))
                    db.db.commit()

                    # Actualizar total de la venta
                    cursor.execute('SELECT SUM(cantidad * valor_venta_producto) AS total FROM detalle_ventas WHERE ID_venta = %s', (venta_id,))
                    total_venta = cursor.fetchone()['total'] or 0
                    cursor.execute('UPDATE ventas SET total = %s WHERE ID_venta = %s', (total_venta, venta_id))
                    db.db.commit()

                    # Actualizar stock del producto
                    new_stock = producto['stock'] - cantidad
                    cursor.execute('UPDATE producto SET stock = %s WHERE ID_producto = %s', (new_stock, id_producto))
                    db.db.commit()

                    flash('Producto agregado a la venta')
                except mysql.connector.Error as err:
                    flash(f"Error al agregar producto a la venta: {err}")
            else:
                flash('No hay suficiente stock')
        else:
            flash('error')

    # Obtener categorías y productos para mostrar en la página
    cursor.execute('SELECT * FROM categoria_producto')
    categorias = cursor.fetchall()

    query = 'SELECT * FROM producto'
    if selected_categoria_id:
        query += ' WHERE ID_categoria_producto = %s'
        cursor.execute(query, (selected_categoria_id,))
    elif search_query:
        query += ' WHERE nombre_producto LIKE %s'
        cursor.execute(query, ('%' + search_query + '%',))
    else:
        cursor.execute(query)

    productos = cursor.fetchall()

    # Obtener detalles de la venta actual
    detalles_venta = []
    total_venta = 0
    if venta_id:
        cursor.execute('SELECT dv.ID_detalle_venta, dv.ID_venta, p.nombre_producto, dv.cantidad, dv.valor_venta_producto '
                    'FROM detalle_ventas dv '
                    'JOIN producto p ON dv.ID_producto = p.ID_producto '
                    'WHERE dv.ID_venta = %s', (venta_id,))
        detalles_venta = cursor.fetchall()

        cursor.execute('SELECT total FROM ventas WHERE ID_venta = %s', (venta_id,))
        total_venta = cursor.fetchone()['total'] or 0

    cursor.close()

    return render_template('ventas.html', categorias=categorias, productos=productos, detalles_venta=detalles_venta, selected_categoria_id=selected_categoria_id, search_query=search_query, total_venta=total_venta, title='usuarios')

@app.route('/ventas/editar/<int:id>', methods=['POST'])
def editar_producto(id):
    if request.method == 'POST':
        nueva_cantidad = int(request.form['cantidad'])

        cursor = db.db.cursor(dictionary=True)
        cursor.execute('SELECT * FROM detalle_ventas WHERE ID_detalle_venta = %s', (id,))
        detalle_venta = cursor.fetchone()

        if detalle_venta:
            id_producto = detalle_venta['ID_producto']
            cursor.execute('SELECT * FROM producto WHERE ID_producto = %s', (id_producto,))
            producto = cursor.fetchone()

            if producto:
                diferencia = nueva_cantidad - detalle_venta['cantidad']
                nuevo_stock = producto['stock'] - diferencia
                cursor.execute('UPDATE producto SET stock = %s WHERE ID_producto = %s', (nuevo_stock, id_producto))
                db.db.commit()

                cursor.execute('UPDATE detalle_ventas SET cantidad = %s WHERE ID_detalle_venta = %s', (nueva_cantidad, id))
                db.db.commit()

                # Actualizar total de la venta
                cursor.execute('SELECT SUM(cantidad * valor_venta_producto) AS total FROM detalle_ventas WHERE ID_venta = %s', (detalle_venta['ID_venta'],))
                total_venta = cursor.fetchone()['total'] or 0
                cursor.execute('UPDATE ventas SET total = %s WHERE ID_venta = %s', (total_venta, detalle_venta['ID_venta']))
                db.db.commit()

                flash('Producto editado correctamente')

    return redirect(url_for('ventas'))

@app.route('/ventas/eliminar/<int:id>', methods=['POST'])
def eliminar_producto(id):
    cursor = db.db.cursor(dictionary=True)
    cursor.execute('SELECT * FROM detalle_ventas WHERE ID_detalle_venta = %s', (id,))
    detalle_venta = cursor.fetchone()

    if detalle_venta:
        id_producto = detalle_venta['ID_producto']
        cantidad = detalle_venta['cantidad']
        cursor.execute('SELECT * FROM producto WHERE ID_producto = %s', (id_producto,))
        producto = cursor.fetchone()

        if producto:
            nuevo_stock = producto['stock'] + cantidad
            cursor.execute('UPDATE producto SET stock = %s WHERE ID_producto = %s', (nuevo_stock, id_producto))
            db.db.commit()

            cursor.execute('DELETE FROM detalle_ventas WHERE ID_detalle_venta = %s', (id,))
            db.db.commit()

            # Actualizar total de la venta
            cursor.execute('SELECT SUM(cantidad * valor_venta_producto) AS total FROM detalle_ventas WHERE ID_venta = %s', (detalle_venta['ID_venta'],))
            total_venta = cursor.fetchone()['total'] or 0
            cursor.execute('UPDATE ventas SET total = %s WHERE ID_venta = %s', (total_venta, detalle_venta['ID_venta']))
            db.db.commit()

            flash('Producto eliminado correctamente')

    return redirect(url_for('ventas'))

@app.route('/ventas/cancelar', methods=['POST'])
def cancelar_venta():
    venta_id = session.get('venta_id', None)

    if venta_id:
        cursor = db.db.cursor(dictionary=True)
        cursor.execute('SELECT * FROM detalle_ventas WHERE ID_venta = %s', (venta_id,))
        detalles_venta = cursor.fetchall()

        for detalle in detalles_venta:
            id_producto = detalle['ID_producto']
            cantidad = detalle['cantidad']
            cursor.execute('SELECT * FROM producto WHERE ID_producto = %s', (id_producto,))
            producto = cursor.fetchone()

            if producto:
                nuevo_stock = producto['stock'] + cantidad
                cursor.execute('UPDATE producto SET stock = %s WHERE ID_producto = %s', (nuevo_stock, id_producto))
                db.db.commit()

        cursor.execute('DELETE FROM detalle_ventas WHERE ID_venta = %s', (venta_id,))
        cursor.execute('DELETE FROM ventas WHERE ID_venta = %s', (venta_id,))
        db.db.commit()

        session.pop('venta_id', None)

        flash('Venta cancelada correctamente')

    return redirect(url_for('ventas'))

@app.route('/ventas/finalizar', methods=['POST'])
def finalizar_venta():
    venta_id = session.get('venta_id', None)

    if venta_id:
        cursor = db.db.cursor(dictionary=True)
        
        # Asegurarse de que el total se ha calculado y actualizado correctamente
        cursor.execute('SELECT SUM(cantidad * valor_venta_producto) AS total FROM detalle_ventas WHERE ID_venta = %s', (venta_id,))
        total_venta = cursor.fetchone()['total'] or 0
        cursor.execute('UPDATE ventas SET total = %s WHERE ID_venta = %s', (total_venta, venta_id))
        db.db.commit()

        session.pop('venta_id', None)

        flash('Venta finalizada correctamente')

    return redirect(url_for('ventas'))

#compras

# Función para obtener la lista de proveedores
def obtener_proveedores():
    try:
        connection = db.db
        cursor = connection.cursor(dictionary=True)
        query = "SELECT * FROM proveedores"
        cursor.execute(query)
        proveedores = cursor.fetchall()
        return proveedores
    except Exception as e:
        print(f"Error al obtener proveedores: {str(e)}")
        return []

# Función para obtener la lista de productos filtrados por proveedor
def obtener_productos_filtrados(proveedor_id):
    try:
        connection = db.db
        cursor = connection.cursor(dictionary=True)
        if proveedor_id:
            query = "SELECT * FROM producto WHERE ID_proveedor = %s ORDER BY nombre_producto ASC"
            cursor.execute(query, (proveedor_id,))
        else:
            query = "SELECT * FROM producto ORDER BY nombre_producto ASC"
            cursor.execute(query)
        productos = cursor.fetchall()
        return productos
    except Exception as e:
        print(f"Error al obtener productos filtrados: {str(e)}")
        return []

# Función para buscar productos por nombre
def buscar_productos(nombre_producto):
    try:
        connection = db.db
        cursor = connection.cursor(dictionary=True)
        query = "SELECT * FROM producto WHERE nombre_producto LIKE %s ORDER BY nombre_producto ASC"
        cursor.execute(query, ('%' + nombre_producto + '%',))
        productos = cursor.fetchall()
        return productos
    except Exception as e:
        print(f"Error al buscar productos: {str(e)}")
        return []

@app.route('/compras', methods=['GET', 'POST'])
def compras():
    
    cursor = db.db.cursor(dictionary=True)

    if request.method == 'POST':
        if 'proveedor_id' in request.form:
            proveedor_id = request.form['proveedor_id']
            productos = obtener_productos_filtrados(proveedor_id)
        elif 'buscar' in request.form:
            nombre_producto = request.form['buscar']
            productos = buscar_productos(nombre_producto)
        else:
            flash('Acción no válida.', 'danger')
            return redirect(url_for('compras'))
    else:
        productos = obtener_productos_filtrados(None)  # Obtener todos los productos al cargar la página
    
    proveedores = obtener_proveedores()
    total_compra = calcular_total_compra(session.get('carrito', []))
    
    return render_template('compras.html', proveedores=proveedores, productos=productos, total_compra=total_compra, carrito=session.get('carrito', []), title='compras')

# Función para agregar al carrito
@app.route('/agregar_al_carrito/<int:producto_id>', methods=['POST'])
def agregar_al_carrito(producto_id):
    cantidad = int(request.form['cantidad'])
    try:
        connection = db.db
        cursor = connection.cursor(dictionary=True)
        query = "SELECT * FROM producto WHERE ID_producto = %s"
        cursor.execute(query, (producto_id,))
        producto = cursor.fetchone()
        
        if producto:
            precio_total = producto['valor_producto'] * cantidad
            producto_carrito = {
                'ID_producto': producto['ID_producto'],
                'nombre_producto': producto['nombre_producto'],
                'cantidad': cantidad,
                'precio_total': precio_total
            }
            carrito = session.get('carrito', [])
            carrito.append(producto_carrito)
            session['carrito'] = carrito
            flash('Producto añadido al carrito.', 'success')
        else:
            flash('Producto no encontrado.', 'danger')
    except Exception as e:
        flash(f'Error al agregar producto al carrito: {str(e)}', 'danger')
    
    return redirect(url_for('compras'))


# Función para eliminar un producto del carrito de compras
@app.route('/eliminar_del_carrito/<int:index>', methods=['POST'])
def eliminar_del_carrito(index):
    carrito = session.get('carrito', [])
    if 0 <= index < len(carrito):
        del carrito[index]
        session['carrito'] = carrito
        flash('Producto eliminado del carrito.', 'success')
    else:
        flash('Índice de carrito no válido.', 'danger')
    
    return redirect(url_for('compras'))

# Función para calcular el total de la compra
def calcular_total_compra(carrito):
    total = 0
    for producto in carrito:
        total += float(producto['precio_total'])  # Asegúrate de que producto['precio_total'] sea numérico
    return total

# Función para cancelar la compra (vaciar carrito)
@app.route('/cancelar_compra', methods=['POST'])
def cancelar_compra():
    session.pop('carrito', None)
    flash('Compra cancelada.', 'info')
    return redirect(url_for('compras'))


@app.route('/finalizar_compra', methods=['POST'])
def finalizar_compra():
    carrito = session.get('carrito', [])
    if carrito:
        try:
            connection = db.db
            cursor = db.get_db_cursor()

            fecha_actual = date.today()
            insert_compra_query = "INSERT INTO compras (fecha_compra, valor_compra) VALUES (%s, %s)"
            cursor.execute(insert_compra_query, (fecha_actual, calcular_total_compra(carrito)))
            connection.commit()

            id_compra = cursor.lastrowid
            for item in carrito:
                producto_id = item.get('ID_producto')
                if producto_id is None:
                    flash(f'El producto en el carrito no tiene un ID válido.', 'danger')
                    continue

                cantidad = item['cantidad']
                precio_total = item['precio_total']

                obtener_proveedor_query = "SELECT ID_proveedor FROM producto WHERE ID_producto = %s"
                cursor.execute(obtener_proveedor_query, (producto_id,))
                proveedor_id_result = cursor.fetchone()

                if proveedor_id_result is None:
                    flash(f'No se encontró proveedor para el producto con ID {producto_id}.', 'warning')
                    continue
                
                proveedor_id = proveedor_id_result['ID_proveedor']
                valor_compra_producto = float(precio_total) * cantidad

                insert_detalle_query = "INSERT INTO detalle_compra (ID_compra, ID_producto, cantidad, ID_proveedor, valor_compra_producto) VALUES (%s, %s, %s, %s, %s)"
                cursor.execute(insert_detalle_query, (id_compra, producto_id, cantidad, proveedor_id, valor_compra_producto))
                connection.commit()

                # Actualizar el stock
                actualizar_stock_query = "UPDATE producto SET stock = stock + %s WHERE ID_producto = %s"
                cursor.execute(actualizar_stock_query, (cantidad, producto_id))
                connection.commit()

            session.pop('carrito', None)
            flash('Compra finalizada correctamente.', 'success')
        
        except Exception as e:
            flash(f'Error al finalizar compra: {str(e)}', 'danger')
            connection.rollback()
            return redirect(url_for('compras'))
        
        finally:
            cursor.close()

    else:
        flash('No hay productos en el carrito para finalizar la compra.', 'warning')

    return redirect(url_for('compras'))


#DEVOLUCIONES

@app.route('/reembolsos', methods=['GET', 'POST'])
def reembolsos():
    cursor = db.get_db_cursor()
    
    # Obtener categorías
    cursor.execute("SELECT * FROM categoria_producto")
    categorias = cursor.fetchall()
    
    # Obtener productos
    cursor.execute("SELECT * FROM producto")
    productos = cursor.fetchall()
    
    if 'carrito_devoluciones' not in session:
        session['carrito_devoluciones'] = []
    
    if request.method == 'POST':
        id_producto = int(request.form['producto'])
        cantidad = int(request.form['cantidad'])
        descripcion = request.form['descripcion']
        retorno_inventario = request.form['retorno_inventario'] == 'true'

        cursor.execute("SELECT * FROM producto WHERE ID_producto = %s", (id_producto,))
        producto = cursor.fetchone()

        if producto:
            session['carrito_devoluciones'].append({
                'ID_producto': producto['ID_producto'],
                'nombre_producto': producto['nombre_producto'],
                'cantidad': cantidad,
                'descripcion': descripcion,
                'retorno_inventario': retorno_inventario
            })
            session.modified = True
            flash('Producto añadido al carrito de devoluciones', 'success')

    carrito_devoluciones = session['carrito_devoluciones']
    return render_template('reembolsos.html', categorias=categorias, productos=productos, carrito_devoluciones=carrito_devoluciones, title='Reembolsos')

@app.route('/eliminar_producto_reembolso/<int:index>', methods=['POST'])
def eliminar_producto_reembolso(index):
    if 'carrito_devoluciones' in session:
        session['carrito_devoluciones'].pop(index)
        session.modified = True
        flash('Producto eliminado del carrito de devoluciones', 'success')
    return redirect(url_for('reembolsos'))



@app.route('/cancelar_reembolso')
def cancelar_reembolso():
    session.pop('carrito_devoluciones', None)
    flash('Reembolso cancelado', 'info')
    return redirect(url_for('reembolsos'))



@app.route('/finalizar_reembolso')
def finalizar_reembolso():
    cursor = db.get_db_cursor()
    
    try:
        # Iniciar transacción
        db.db.begin()

        # Crear registro de devolución
        cursor.execute("INSERT INTO devoluciones (fecha_devolucion) VALUES (%s)", (date.today(),))
        id_devolucion = cursor.lastrowid

        # Verificar si 'carrito_devoluciones' está en la sesión
        if 'carrito_devoluciones' not in session or not session['carrito_devoluciones']:
            raise ValueError('El carrito de devoluciones está vacío o no existe')

        # Insertar detalles de devolución y actualizar stock si es necesario
        for item in session['carrito_devoluciones']:
            cursor.execute(
                "INSERT INTO detalle_devolucion (ID_devolucion, ID_producto, descripcion, cantidad, retorno_inventario) VALUES (%s, %s, %s, %s, %s)",
                (id_devolucion, item['ID_producto'], item['descripcion'], item['cantidad'], item['retorno_inventario'])
            )
            if item['retorno_inventario']:
                cursor.execute("UPDATE producto SET stock = stock + %s WHERE ID_producto = %s", (item['cantidad'], item['ID_producto']))
        
        # Commit the transaction
        db.db.commit()
        session.pop('carrito_devoluciones', None)
        flash('Reembolso finalizado', 'success')

    except Exception as e:
        # Rollback the transaction in case of error
        db.db.rollback()
        flash('Error al finalizar el reembolso: {}'.format(str(e)), 'danger')
    
    return redirect(url_for('reembolsos'))





#DOWNLOAD REGISTROS


# Function to fetch data from the database
def fetch_data(query):
    cursor = db.get_db_cursor()
    cursor.execute(query)
    return cursor.fetchall()

# Function to fetch sales data
def fetch_sales_data(start_date, end_date):
    query_detail = f"""
    SELECT ventas.ID_venta, usuarios.Nombre AS Nombre_usuario, producto.nombre_producto, detalle_ventas.valor_venta_producto, 
    detalle_ventas.cantidad, (detalle_ventas.valor_venta_producto * detalle_ventas.cantidad) AS total 
    FROM detalle_ventas 
    JOIN ventas ON detalle_ventas.ID_venta = ventas.ID_venta 
    JOIN usuarios ON ventas.ID_usuario = usuarios.ID_usuario 
    JOIN producto ON detalle_ventas.ID_producto = producto.ID_producto 
    WHERE ventas.fecha_venta BETWEEN '{start_date}' AND '{end_date}'
    """
    sales_detail = fetch_data(query_detail)

    query_sales = f"""
    SELECT ID_venta, fecha_venta, total FROM ventas WHERE fecha_venta BETWEEN '{start_date}' AND '{end_date}'
    """
    sales = fetch_data(query_sales)

    return sales_detail, sales

# Function to fetch purchases data
def fetch_purchases_data(start_date, end_date):
    query_detail = f"""
    SELECT detalle_compra.ID_compra, producto.nombre_producto, detalle_compra.cantidad, proveedores.Nombre_empresa 
    FROM detalle_compra 
    JOIN producto ON detalle_compra.ID_producto = producto.ID_producto 
    JOIN proveedores ON detalle_compra.ID_proveedor = proveedores.ID_proveedor 
    JOIN compras ON detalle_compra.ID_compra = compras.ID_compra 
    WHERE compras.fecha_compra BETWEEN '{start_date}' AND '{end_date}'
    """
    purchase_detail = fetch_data(query_detail)

    query_purchases = f"""
    SELECT ID_compra, fecha_compra FROM compras WHERE fecha_compra BETWEEN '{start_date}' AND '{end_date}'
    """
    purchases = fetch_data(query_purchases)

    return purchase_detail, purchases

# Function to fetch returns data
def fetch_returns_data(start_date, end_date):
    query_detail = f"""
    SELECT detalle_devolucion.ID_devolucion, producto.nombre_producto, detalle_devolucion.descripcion, 
    CASE WHEN detalle_devolucion.retorno_inventario = 1 THEN 'Sí' ELSE 'No' END AS retorno_inventario 
    FROM detalle_devolucion 
    JOIN producto ON detalle_devolucion.ID_producto = producto.ID_producto 
    JOIN devoluciones ON detalle_devolucion.ID_devolucion = devoluciones.ID_devolucion 
    WHERE devoluciones.fecha_devolucion BETWEEN '{start_date}' AND '{end_date}'
    """
    return_detail = fetch_data(query_detail)

    query_returns = f"""
    SELECT ID_devolucion, fecha_devolucion FROM devoluciones WHERE fecha_devolucion BETWEEN '{start_date}' AND '{end_date}'
    """
    returns = fetch_data(query_returns)

    return return_detail, returns

# Function to fetch employees data
def fetch_employees_data():
    query = """
    SELECT usuarios.Nombre, COUNT(ventas.ID_venta) AS num_ventas, SUM(ventas.total) AS total_ventas 
    FROM usuarios 
    JOIN ventas ON usuarios.ID_usuario = ventas.ID_usuario 
    GROUP BY usuarios.Nombre
    """
    return fetch_data(query)

# Function to generate PDF
def generate_pdf(data):
    buffer = io.BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=letter)
    pdf.setTitle("Reporte de Tienda")
    
    # Set margins and page dimensions
    margin_bottom = 50
    margin_left = 50
    page_width, page_height = letter
    
    # Start at the top of the page
    y_position = page_height - margin_bottom

    pdf.drawString(margin_left, y_position, "Reporte de Tienda")
    y_position -= 20  # Move down after title

    if 'sales_detail' in data:
        pdf.drawString(margin_left, y_position, "Ventas - Detalles:")
        y_position -= 20
        for sale in data['sales_detail']:
            if y_position < margin_bottom:  # Check if at the bottom of the page
                pdf.showPage()  # Start a new page if content exceeds
                y_position = page_height - margin_bottom
            
            # Adjust font size based on content
            font_size = 10
            text = f"ID Venta: {sale['ID_venta']} - Usuario: {sale['Nombre_usuario']} - Producto: {sale['nombre_producto']} - Valor Unitario: {sale['valor_venta_producto']} - Cantidad: {sale['cantidad']} - Total: {sale['total']}"
            while pdf.stringWidth(text, font_size) > page_width - 2 * margin_left:
                font_size -= 1
            pdf.setFont("Helvetica", font_size)
            
            pdf.drawString(margin_left, y_position, text)
            y_position -= 20

        y_position -= 20  # Space between sections

        pdf.drawString(margin_left, y_position, "Ventas:")
        y_position -= 20
        for sale in data['sales']:
            if y_position < margin_bottom:  # Check if at the bottom of the page
                pdf.showPage()  # Start a new page if content exceeds
                y_position = page_height - margin_bottom
            
            # Adjust font size based on content
            font_size = 10
            text = f"ID Venta: {sale['ID_venta']} - Fecha: {sale['fecha_venta']} - Total: {sale['total']}"
            while pdf.stringWidth(text, font_size) > page_width - 2 * margin_left:
                font_size -= 1
            pdf.setFont("Helvetica", font_size)
            
            pdf.drawString(margin_left, y_position, text)
            y_position -= 20

    # Add more sections similarly for other data

    pdf.save()
    buffer.seek(0)
    return buffer

# Function to generate Excel
def generate_excel(data):
    output = io.BytesIO()

    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        if 'sales_detail' in data:
            df_sales_detail = pd.DataFrame(data['sales_detail'])
            df_sales_detail.to_excel(writer, sheet_name='Detalle Ventas', index=False)

        if 'sales' in data:
            df_sales = pd.DataFrame(data['sales'])
            df_sales.to_excel(writer, sheet_name='Ventas', index=False)

        if 'purchase_detail' in data:
            df_purchase_detail = pd.DataFrame(data['purchase_detail'])
            df_purchase_detail.to_excel(writer, sheet_name='Detalle Compras', index=False)

        if 'purchases' in data:
            df_purchases = pd.DataFrame(data['purchases'])
            df_purchases.to_excel(writer, sheet_name='Compras', index=False)

        if 'return_detail' in data:
            df_return_detail = pd.DataFrame(data['return_detail'])
            df_return_detail.to_excel(writer, sheet_name='Detalle Devoluciones', index=False)

        if 'returns' in data:
            df_returns = pd.DataFrame(data['returns'])
            df_returns.to_excel(writer, sheet_name='Devoluciones', index=False)

        if 'employees' in data:
            df_employees = pd.DataFrame(data['employees'])
            df_employees.to_excel(writer, sheet_name='Empleados', index=False)

    output.seek(0)  # Ensure the buffer's cursor is at the beginning before returning
    return output

# Function to generate Word
def generate_word(data):
    doc = Document()
    doc.add_heading('Reporte de Tienda', 0)
    doc.add_paragraph('Este es el reporte generado para la tienda.')

    if 'sales_detail' in data:
        doc.add_heading('Ventas - Detalles', level=1)
        for sale in data['sales_detail']:
            doc.add_paragraph(
                f"ID Venta: {sale['ID_venta']} - Usuario: {sale['Nombre_usuario']} - Producto: {sale['nombre_producto']} - Valor Unitario: {sale['valor_venta_producto']} - Cantidad: {sale['cantidad']} - Total: {sale['total']}"
            )

        doc.add_heading('Ventas', level=1)
        for sale in data['sales']:
            doc.add_paragraph(
                f"ID Venta: {sale['ID_venta']} - Fecha: {sale['fecha_venta']} - Total: {sale['total']}"
            )

    if 'purchase_detail' in data:
        doc.add_heading('Compras - Detalles', level=1)
        for purchase in data['purchase_detail']:
            doc.add_paragraph(
                f"ID Compra: {purchase['ID_compra']} - Producto: {purchase['nombre_producto']} - Cantidad: {purchase['cantidad']} - Proveedor: {purchase['Nombre_empresa']}"
            )

        doc.add_heading('Compras', level=1)
        for purchase in data['purchases']:
            doc.add_paragraph(
                f"ID Compra: {purchase['ID_compra']} - Fecha: {purchase['fecha_compra']}"
            )

    if 'return_detail' in data:
        doc.add_heading('Devoluciones - Detalles', level=1)
        for return_ in data['return_detail']:
            doc.add_paragraph(
                f"ID Devolución: {return_['ID_devolucion']} - Producto: {return_['nombre_producto']} - Descripción: {return_['descripcion']} - ¿Retorno a Inventario?: {return_['retorno_inventario']}"
            )

        doc.add_heading('Devoluciones', level=1)
        for return_ in data['returns']:
            doc.add_paragraph(
                f"ID Devolución: {return_['ID_devolucion']} - Fecha: {return_['fecha_devolucion']}"
            )

    if 'employees' in data:
        doc.add_heading('Empleados', level=1)
        for employee in data['employees']:
            doc.add_paragraph(
                f"Nombre: {employee['Nombre']} - Número de Ventas: {employee['num_ventas']} - Total Ventas: {employee['total_ventas']}"
            )

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output

@app.route('/generate-report', methods=['POST'])
def generate_report():
    request_data = request.get_json()
    format = request_data['format']
    start_date = request_data['start_date']
    end_date = request_data['end_date']
    selected_records = request_data['selected_records']

    data = {}
    if 'ventas' in selected_records:
        data['sales_detail'], data['sales'] = fetch_sales_data(start_date, end_date)
    if 'compras' in selected_records:
        data['purchase_detail'], data['purchases'] = fetch_purchases_data(start_date, end_date)
    if 'devoluciones' in selected_records:
        data['return_detail'], data['returns'] = fetch_returns_data(start_date, end_date)
    if 'empleados' in selected_records:
        data['employees'] = fetch_employees_data()

    if format == 'pdf':
        buffer = generate_pdf(data)
        return send_file(buffer, as_attachment=True, download_name="reporte_tienda.pdf", mimetype='application/pdf')
    elif format == 'xlsx':
        buffer = generate_excel(data)
        return send_file(buffer, as_attachment=True, download_name="reporte_tienda.xlsx", mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    elif format == 'docx':
        buffer = generate_word(data)
        return send_file(buffer, as_attachment=True, download_name="reporte_tienda.docx", mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')




#CATEGORIAS

@app.route('/add_category', methods=['POST'])
def add_category():
    category_name = request.form['category_name']
    cursor = get_db_cursor()
    cursor.execute("INSERT INTO categoria_producto (nombre_categoria) VALUES (%s)", (category_name,))
    database.commit()  # Commit on the actual database connection object
    cursor.close()
    return redirect(url_for('inventario'))

@app.route('/search_category')
def search_category():
    category_name = request.args.get('name', '')
    cursor = get_db_cursor()
    cursor.execute("SELECT * FROM categoria_producto WHERE nombre_categoria LIKE %s", (f"%{category_name}%",))
    categories = cursor.fetchall()
    cursor.close()
    return jsonify(categories)

@app.route('/update_category', methods=['POST'])
def update_category():
    category_id = request.form['id']
    category_name = request.form['category_name']
    cursor = get_db_cursor()
    cursor.execute("UPDATE categoria_producto SET nombre_categoria = %s WHERE ID_categoria_producto = %s", (category_name, category_id))
    database.commit()  # Commit on the actual database connection object
    cursor.close()
    return redirect(url_for('inventario'))



@app.route('/historial_ventas')
def historial_ventas():
    cursor = get_db_cursor()
    # Fetch sales data
    cursor.execute('''
        SELECT ID_venta, fecha_venta, total
        FROM ventas
    ''')
    ventas = cursor.fetchall()

    for venta in ventas:
        cursor.execute('''
            SELECT 
                usuarios.Nombre as nombre_usuario,
                producto.nombre_producto,
                detalle_ventas.cantidad,
                producto.valor_producto,
                (detalle_ventas.cantidad * producto.valor_producto) as total_precio
            FROM detalle_ventas
            JOIN ventas ON ventas.ID_venta = detalle_ventas.ID_venta
            JOIN usuarios ON usuarios.ID_usuario = ventas.ID_usuario
            JOIN producto ON producto.ID_producto = detalle_ventas.ID_producto
            WHERE detalle_ventas.ID_venta = %s
        ''', (venta['ID_venta'],))
        venta['detalles'] = cursor.fetchall()
    
    cursor.close()
    return render_template('historial_ventas.html', ventas=ventas)


@app.route('/historial_compras')
def historial_compras():
    cursor = get_db_cursor()
    # Fetch purchase data
    cursor.execute('''
        SELECT ID_compra, fecha_compra, valor_compra
        FROM compras
    ''')
    compras = cursor.fetchall()

    for compra in compras:
        cursor.execute('''
            SELECT 
                producto.nombre_producto,
                detalle_compra.cantidad,
                detalle_compra.valor_compra_producto
            FROM detalle_compra
            JOIN producto ON producto.ID_producto = detalle_compra.ID_producto
            WHERE detalle_compra.ID_compra = %s
        ''', (compra['ID_compra'],))
        compra['detalles'] = cursor.fetchall()
    
    cursor.close()
    return render_template('historial_compras.html', compras=compras)


if __name__ == '__main__':
    app.run(debug=True, port=5000)


